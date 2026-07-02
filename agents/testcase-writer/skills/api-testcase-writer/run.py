# -*- coding: utf-8 -*-
"""
Skill: api-testcase-writer

API测试用例编写器：根据需求材料生成测试用例Excel。
机械操作：init-materials, generate-excel, convert-interface-doc, list-channels, list-interface-docs
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import sys
from datetime import datetime
from typing import Any, Dict, List, Optional

import openpyxl

# ── 路径常量（TestScheduler 独立资源） ──
SKILL_DIR = os.path.dirname(os.path.abspath(__file__))
TESTCASE_DIR = "D:/TestScheduler/memory/testcases"
REPORT_DIR = "D:/TestScheduler/memory/test_results/reports"
TEMPLATE_FILE = os.path.join(TESTCASE_DIR, "template.xlsx")
CHANNELS_DIR = "D:/TestScheduler/memory/api_channels"
CHANNEL_RULES_DIR = "D:/TestScheduler/memory/api_channels_rules"
REQUIREMENT_MATERIALS_DIR = "D:/TestScheduler/memory/requirement_materials"
INTERFACE_DOCS_DIR = "D:/接口文档合集"
# 接口文档 MD 缓存目录（TestScheduler 独立管理）
INTERFACE_DOC_CACHE_DIR = "D:/TestScheduler/memory/interface_docs"
CACHE_INDEX_FILE = os.path.join(INTERFACE_DOC_CACHE_DIR, "_cache_index.json")

# ── Excel 列映射（与 api-testcase-executor 对齐） ────────
COL_MAP = {
    "story_id": 0,
    "node_path": 1,
    "name": 2,
    "summary": 3,
    "precondition": 4,
    "steps": 5,
    "expected": 6,
    "actual": 7,
    "priority": 8,
    "author": 9,
    "type": 10,
    "need_regression": 11,
    "app": 12,
    "tags": 13,
}

# 列名（用于无模板时创建表头）
COL_NAMES = [
    "STORY编号", "节点路径", "测试用例名称", "摘要", "前置条件",
    "执行步骤", "预期结果", "实际结果", "优先级", "作者", "类型", "是否需要回归", "应用", "标签",
]


# ═══════════════════════════════════════════════════════
# 1. init-materials
# ═══════════════════════════════════════════════════════

def init_materials(params: dict) -> str:
    """创建需求材料目录和模板txt文件"""
    req_id = params.get("requirement_id", "").strip()
    if not req_id:
        return json.dumps({"ok": False, "error": "缺少 requirement_id"}, ensure_ascii=False)

    # 目录结构: requirement_materials/{req_id}/{req_id}.txt + 其他文件
    req_dir = os.path.join(REQUIREMENT_MATERIALS_DIR, req_id)
    os.makedirs(req_dir, exist_ok=True)

    txt_path = os.path.join(req_dir, f"{req_id}.txt")

    if os.path.exists(txt_path):
        with open(txt_path, "r", encoding="utf-8") as f:
            existing_content = f.read()
        # 同时列出目录下其他文件
        other_files = [
            f for f in os.listdir(req_dir)
            if f != f"{req_id}.txt" and not f.startswith("~$")
        ]
        return json.dumps({
            "ok": True,
            "message": f"材料目录已存在: {req_dir}",
            "dir_path": req_dir,
            "txt_path": txt_path,
            "txt_content": existing_content,
            "other_files": other_files,
        }, ensure_ascii=False, indent=2)

    template_content = f"""需求文档:
接口文档:
开发设计文档:
测试人员补充:
代码目录:
"""

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(template_content)

    return json.dumps({
        "ok": True,
        "message": f"已创建材料目录: {req_dir}，请填写 {req_id}.txt 或直接放入文件",
        "dir_path": req_dir,
        "txt_path": txt_path,
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════
# 2. generate-excel
# ═══════════════════════════════════════════════════════

def generate_excel(params: dict) -> str:
    """将用例JSON数组写入Excel文件"""
    cases_json = params.get("cases_json", "")
    output_name = params.get("output_name", "").strip()

    if not cases_json:
        return json.dumps({"ok": False, "error": "缺少 cases_json（用例JSON数组）"}, ensure_ascii=False)

    try:
        if isinstance(cases_json, str):
            cases = json.loads(cases_json)
        else:
            cases = cases_json
    except json.JSONDecodeError as e:
        return json.dumps({"ok": False, "error": f"cases_json 解析失败: {e}"}, ensure_ascii=False)

    if not isinstance(cases, list) or not cases:
        return json.dumps({"ok": False, "error": "cases_json 必须是非空数组"}, ensure_ascii=False)

    # 校验每个用例必须有 name, steps, expected
    for i, case in enumerate(cases):
        if not case.get("name"):
            return json.dumps({"ok": False, "error": f"第{i+1}个用例缺少 name"}, ensure_ascii=False)
        if not case.get("steps"):
            return json.dumps({"ok": False, "error": f"第{i+1}个用例缺少 steps"}, ensure_ascii=False)
        if not case.get("expected"):
            return json.dumps({"ok": False, "error": f"第{i+1}个用例缺少 expected"}, ensure_ascii=False)

    # 输出文件名
    if not output_name:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_name = f"testcases_{timestamp}"

    output_path = os.path.join(TESTCASE_DIR, f"{output_name}_testcases.xlsx")

    # 加载模板或创建新工作簿
    if os.path.exists(TEMPLATE_FILE):
        wb = openpyxl.load_workbook(TEMPLATE_FILE)
        ws = wb.active
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "测试用例"
        # 写表头
        for col_idx, col_name in enumerate(COL_NAMES):
            ws.cell(row=1, column=col_idx + 1, value=col_name)

    # 找到数据起始行（跳过表头和已有数据）
    start_row = 2
    for row in ws.iter_rows(min_row=2, max_col=1, values_only=True):
        if row[0] is not None:
            start_row += 1
        else:
            break

    # 写入用例数据
    for case in cases:
        row_data = [
            case.get("story_id", ""),
            case.get("node_path", ""),
            case.get("name", ""),
            case.get("summary", ""),
            case.get("precondition", ""),
            case.get("steps", ""),
            case.get("expected", ""),
            case.get("actual", ""),
            case.get("priority", ""),
            case.get("author", ""),
            case.get("type", "功能测试"),
            case.get("need_regression", "否"),
            case.get("app", ""),
            case.get("tags", ""),
        ]

        for col_idx, value in enumerate(row_data):
            ws.cell(row=start_row, column=col_idx + 1, value=str(value) if value else "")

        start_row += 1

    wb.save(output_path)

    return json.dumps({
        "ok": True,
        "message": f"已生成 {len(cases)} 条用例",
        "file_path": output_path,
        "total_cases": len(cases),
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════
# 3. convert-interface-doc
# ═══════════════════════════════════════════════════════

def _compute_file_hash(file_path: str) -> str:
    """计算文件SHA256哈希"""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _load_cache_index() -> dict:
    """加载缓存索引"""
    if os.path.exists(CACHE_INDEX_FILE):
        try:
            with open(CACHE_INDEX_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def _save_cache_index(index: dict):
    """保存缓存索引"""
    os.makedirs(INTERFACE_DOC_CACHE_DIR, exist_ok=True)
    with open(CACHE_INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False, indent=2)


def _xlsx_to_markdown(file_path: str) -> str:
    """将xlsx文件转为markdown"""
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    sections = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sections.append(f"\n## {sheet_name}\n")

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # 表头
        header = [str(c).strip() if c is not None else "" for c in rows[0]]
        if not any(header):
            continue

        sections.append("| " + " | ".join(header) + " |")
        sections.append("| " + " | ".join(["---"] * len(header)) + " |")

        # 数据行
        for row in rows[1:]:
            cells = []
            for cell in row:
                val = str(cell).strip() if cell is not None else ""
                # 转义markdown表格中的管道符
                val = val.replace("|", "\\|")
                # 截断过长的内容
                if len(val) > 200:
                    val = val[:200] + "..."
                cells.append(val)
            sections.append("| " + " | ".join(cells) + " |")

    wb.close()
    return "\n".join(sections)


def _doc_to_markdown(file_path: str) -> str:
    """将doc/docx文件转为markdown"""
    # 优先尝试 python-docx
    try:
        from docx import Document
        doc = Document(file_path)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""
            if "Heading 1" in style:
                lines.append(f"# {text}")
            elif "Heading 2" in style:
                lines.append(f"## {text}")
            elif "Heading 3" in style:
                lines.append(f"### {text}")
            else:
                lines.append(text)
        # 表格
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        return "\n".join(lines)
    except ImportError:
        pass

    # 回退到 pandoc
    try:
        import subprocess
        result = subprocess.run(
            ["pandoc", file_path, "-t", "markdown", "--wrap=none"],
            capture_output=True, text=True, timeout=60, encoding="utf-8",
        )
        if result.returncode == 0:
            return result.stdout
        return f"pandoc 转换失败: {result.stderr}"
    except FileNotFoundError:
        return "python-docx 和 pandoc 均不可用，无法转换 doc 文件"
    except Exception as e:
        return f"转换异常: {e}"


def _pdf_to_markdown(file_path: str) -> str:
    """将pdf文件转为markdown"""
    try:
        import subprocess
        result = subprocess.run(
            ["pandoc", file_path, "-t", "markdown", "--wrap=none"],
            capture_output=True, text=True, timeout=60, encoding="utf-8",
        )
        if result.returncode == 0:
            return result.stdout
        return f"pandoc 转换失败: {result.stderr}"
    except FileNotFoundError:
        # 回退到 pymupdf
        try:
            import fitz
            doc = fitz.open(file_path)
            lines = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    lines.append(text.strip())
            doc.close()
            return "\n\n".join(lines)
        except ImportError:
            return "pandoc 和 pymupdf 均不可用，无法转换 pdf 文件"
    except Exception as e:
        return f"转换异常: {e}"


def convert_interface_doc(params: dict) -> str:
    """将接口文档转为markdown并缓存"""
    source_file = params.get("source_file", "").strip()
    channel_name = params.get("channel_name", "").strip()

    if not source_file:
        return json.dumps({"ok": False, "error": "缺少 source_file"}, ensure_ascii=False)

    if not os.path.exists(source_file):
        return json.dumps({"ok": False, "error": f"文件不存在: {source_file}"}, ensure_ascii=False)

    # 计算哈希
    file_hash = _compute_file_hash(source_file)

    # 检查缓存
    cache_index = _load_cache_index()
    cache_entry = cache_index.get(source_file, {})

    if cache_entry.get("hash") == file_hash:
        cached_md = cache_entry.get("md_path", "")
        if cached_md and os.path.exists(cached_md):
            return json.dumps({
                "ok": True,
                "message": "缓存命中，直接返回",
                "md_path": cached_md,
                "from_cache": True,
            }, ensure_ascii=False, indent=2)

    # 转换
    ext = os.path.splitext(source_file)[1].lower()
    if ext == ".xlsx":
        md_content = _xlsx_to_markdown(source_file)
    elif ext in (".doc", ".docx"):
        md_content = _doc_to_markdown(source_file)
    elif ext == ".pdf":
        md_content = _pdf_to_markdown(source_file)
    else:
        return json.dumps({"ok": False, "error": f"不支持的文件格式: {ext}"}, ensure_ascii=False)

    # 写缓存
    os.makedirs(INTERFACE_DOC_CACHE_DIR, exist_ok=True)

    if not channel_name:
        channel_name = os.path.splitext(os.path.basename(source_file))[0]

    md_path = os.path.join(INTERFACE_DOC_CACHE_DIR, f"{channel_name}.md")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# {channel_name} 接口文档\n\n")
        f.write(f"> 源文件: {source_file}\n")
        f.write(f"> 转换时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"> 源文件hash: {file_hash}\n\n")
        f.write(md_content)

    # 更新索引
    cache_index[source_file] = {
        "hash": file_hash,
        "md_path": md_path,
        "converted_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    _save_cache_index(cache_index)

    return json.dumps({
        "ok": True,
        "message": "转换完成并已缓存",
        "md_path": md_path,
        "from_cache": False,
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════
# 4. list-channels
# ═══════════════════════════════════════════════════════

def list_channels(params: dict) -> str:
    """列出可用渠道JSON"""
    if not os.path.exists(CHANNELS_DIR):
        return json.dumps({"ok": True, "channels": [], "total": 0}, ensure_ascii=False)

    channels = []
    for fname in sorted(os.listdir(CHANNELS_DIR)):
        if not fname.endswith(".json"):
            continue

        fpath = os.path.join(CHANNELS_DIR, fname)
        try:
            with open(fpath, "r", encoding="utf-8") as f:
                data = json.load(f)

            ch = data.get("channel", {})
            interfaces = data.get("interfaces", [])
            channels.append({
                "partner_code": ch.get("partner_code", ""),
                "partner_desc": ch.get("partner_desc", ""),
                "app_id": ch.get("app_id", ""),
                "interface_count": len(interfaces),
                "methods": [i.get("method", "") for i in interfaces],
            })
        except Exception:
            channels.append({
                "partner_code": fname.replace(".json", ""),
                "partner_desc": "(读取失败)",
                "interface_count": 0,
                "methods": [],
            })

    return json.dumps({
        "ok": True,
        "total": len(channels),
        "channels": channels,
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════
# 5. list-interface-docs
# ═══════════════════════════════════════════════════════

def list_interface_docs(params: dict) -> str:
    """列出可用接口文档（原始文档 + 已缓存MD文件）"""
    docs = []
    seen_channels = set()

    # ① 优先列出已缓存的 MD 文件
    if os.path.exists(INTERFACE_DOC_CACHE_DIR):
        for fname in sorted(os.listdir(INTERFACE_DOC_CACHE_DIR)):
            if fname.startswith("_") or fname.startswith("~$"):
                continue
            fpath = os.path.join(INTERFACE_DOC_CACHE_DIR, fname)
            if not os.path.isfile(fpath):
                continue
            ext = os.path.splitext(fname)[1].lower()
            if ext != ".md":
                continue
            size = os.path.getsize(fpath)
            channel_name = fname.replace(".md", "")
            seen_channels.add(channel_name)
            docs.append({
                "filename": fname,
                "channel_name": channel_name,
                "extension": ".md",
                "size_kb": round(size / 1024, 1),
                "full_path": fpath,
                "source": "cache",
            })

    # ② 列出原始接口文档（排除已缓存的渠道）
    if os.path.exists(INTERFACE_DOCS_DIR):
        cache_index = _load_cache_index()
        for fname in sorted(os.listdir(INTERFACE_DOCS_DIR)):
            if fname.startswith("~$"):
                continue
            fpath = os.path.join(INTERFACE_DOCS_DIR, fname)
            if not os.path.isfile(fpath):
                continue

            ext = os.path.splitext(fname)[1].lower()
            size = os.path.getsize(fpath)

            # 检查是否已有缓存
            cached = cache_index.get(fpath, {})
            is_cached = bool(cached.get("md_path") and os.path.exists(cached["md_path"]))

            docs.append({
                "filename": fname,
                "extension": ext,
                "size_kb": round(size / 1024, 1),
                "full_path": fpath,
                "source": "original",
                "is_cached": is_cached,
            })

    return json.dumps({
        "ok": True,
        "total": len(docs),
        "cached_count": sum(1 for d in docs if d.get("source") == "cache"),
        "original_count": sum(1 for d in docs if d.get("source") == "original"),
        "docs": docs,
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════
# 5.5. find-interface-doc
# ═══════════════════════════════════════════════════════

def _fuzzy_match_channel(query: str, candidates: list) -> Optional[str]:
    """模糊匹配渠道名，返回最佳匹配的文件路径或None"""
    if not query:
        return None
    query_lower = query.lower().replace(" ", "").replace("_", "").replace("-", "")

    # 精确匹配优先
    for c in candidates:
        c_name = os.path.splitext(os.path.basename(c))[0].lower().replace(" ", "").replace("_", "").replace("-", "")
        if query_lower == c_name:
            return c

    # 包含匹配
    for c in candidates:
        c_name = os.path.splitext(os.path.basename(c))[0].lower().replace(" ", "").replace("_", "").replace("-", "")
        if query_lower in c_name or c_name in query_lower:
            return c

    return None


def find_interface_doc(params: dict) -> str:
    """根据渠道名查找接口文档，优先MD缓存→原始文档→转换→提示用户"""
    channel_name = params.get("channel_name", "").strip()
    if not channel_name:
        return json.dumps({"ok": False, "error": "缺少 channel_name 参数"}, ensure_ascii=False)

    # ① 先在 MD 缓存目录中查找
    md_candidates = []
    if os.path.exists(INTERFACE_DOC_CACHE_DIR):
        for fname in os.listdir(INTERFACE_DOC_CACHE_DIR):
            if fname.endswith(".md") and not fname.startswith("_"):
                md_candidates.append(os.path.join(INTERFACE_DOC_CACHE_DIR, fname))

    md_match = _fuzzy_match_channel(channel_name, md_candidates)
    if md_match:
        return json.dumps({
            "ok": True,
            "found": True,
            "source": "cache",
            "md_path": md_match,
            "channel_name": channel_name,
            "message": f"在缓存中找到渠道 [{channel_name}] 的接口文档 MD",
        }, ensure_ascii=False, indent=2)

    # ② 在原始接口文档目录中查找
    original_candidates = []
    if os.path.exists(INTERFACE_DOCS_DIR):
        for fname in os.listdir(INTERFACE_DOCS_DIR):
            if not fname.startswith("~$") and os.path.isfile(os.path.join(INTERFACE_DOCS_DIR, fname)):
                original_candidates.append(os.path.join(INTERFACE_DOCS_DIR, fname))

    original_match = _fuzzy_match_channel(channel_name, original_candidates)
    if original_match:
        # 找到原始文档 → 自动转换为 MD 并缓存
        ext = os.path.splitext(original_match)[1].lower()
        if ext in (".xlsx", ".xls", ".doc", ".docx", ".pdf"):
            convert_result = convert_interface_doc({
                "source_file": original_match,
                "channel_name": channel_name,
            })
            convert_data = json.loads(convert_result)
            if convert_data.get("ok"):
                return json.dumps({
                    "ok": True,
                    "found": True,
                    "source": "original",
                    "original_path": original_match,
                    "md_path": convert_data.get("md_path"),
                    "message": f"首次识别渠道 [{channel_name}] 的接口文档，已自动转换为 MD 并缓存",
                    "is_first_time": True,
                }, ensure_ascii=False, indent=2)
            else:
                return json.dumps({
                    "ok": True,
                    "found": True,
                    "source": "original",
                    "original_path": original_match,
                    "md_path": None,
                    "message": f"找到原始文档但转换失败: {convert_data.get('error')}",
                    "warning": "转换失败，请手动检查文件格式",
                }, ensure_ascii=False, indent=2)
        else:
            return json.dumps({
                "ok": True,
                "found": True,
                "source": "original",
                "original_path": original_match,
                "md_path": None,
                "message": f"找到原始文档但格式不支持自动转换 ({ext})",
            }, ensure_ascii=False, indent=2)

    # ③ 都没有找到 → 提示用户提供
    return json.dumps({
        "ok": True,
        "found": False,
        "channel_name": channel_name,
        "message": f"未找到渠道 [{channel_name}] 的接口文档。请在 D:/接口文档合集/ 中放入接口文档，或直接在 memory/interface_docs/ 中放入 MD 文件。",
        "suggestion": "请提供接口文档（xlsx/docx/pdf），或直接提供 MD 格式的接口说明",
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════
# 6. scan-materials
# ═══════════════════════════════════════════════════════

def scan_materials(params: dict) -> str:
    """扫描需求材料目录，返回txt内容和目录下所有文件列表"""
    req_id = params.get("requirement_id", "").strip()
    if not req_id:
        return json.dumps({"ok": False, "error": "缺少 requirement_id"}, ensure_ascii=False)

    req_dir = os.path.join(REQUIREMENT_MATERIALS_DIR, req_id)
    if not os.path.exists(req_dir):
        return json.dumps({"ok": False, "error": f"需求目录不存在: {req_dir}，请先调用 init-materials"}, ensure_ascii=False)

    # 读取txt
    txt_path = os.path.join(req_dir, f"{req_id}.txt")
    txt_content = ""
    if os.path.exists(txt_path):
        with open(txt_path, "r", encoding="utf-8") as f:
            txt_content = f.read()

    # 扫描目录下其他文件
    files = []
    for fname in sorted(os.listdir(req_dir)):
        if fname.startswith("~$"):
            continue
        fpath = os.path.join(req_dir, fname)
        if not os.path.isfile(fpath):
            continue

        ext = os.path.splitext(fname)[1].lower()
        size = os.path.getsize(fpath)

        file_info = {
            "filename": fname,
            "extension": ext,
            "size_kb": round(size / 1024, 1),
            "full_path": fpath,
            "is_txt": fname == f"{req_id}.txt",
        }

        # 标记文件类型
        if ext in (".xlsx", ".xls"):
            file_info["type"] = "interface_doc"
        elif ext in (".doc", ".docx"):
            file_info["type"] = "doc"
        elif ext == ".pdf":
            file_info["type"] = "pdf"
        elif ext in (".md",):
            file_info["type"] = "markdown"
        elif ext == ".json":
            file_info["type"] = "json"
        elif ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
            file_info["type"] = "image"
        else:
            file_info["type"] = "other"

        files.append(file_info)

    return json.dumps({
        "ok": True,
        "req_dir": req_dir,
        "txt_path": txt_path,
        "txt_content": txt_content,
        "files": files,
        "total_files": len(files),
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════
# 7. check-existing-testcases
# ═══════════════════════════════════════════════════════

def check_existing_testcases(params: dict) -> str:
    """检查需求是否已有用例文件、测试报告、工作流历史，返回完整结构化结果"""
    req_id = params.get("requirement_id", "").strip()
    if not req_id:
        return json.dumps({"ok": False, "error": "缺少 requirement_id"}, ensure_ascii=False)

    import glob

    # ① 扫描本地用例文件
    tc_pattern = os.path.join(TESTCASE_DIR, f"{req_id}_testcases_*.xlsx")
    tc_matches = sorted(glob.glob(tc_pattern))

    # ② 扫描执行报告
    report_dir = "D:/TestScheduler/memory/test_results/reports"
    rp_patterns = [
        os.path.join(report_dir, f"{req_id}_*.html"),
        os.path.join(report_dir, f"{req_id}_*.md"),
        os.path.join(report_dir, f"{req_id}_*.json"),
    ]
    rp_matches = []
    for pat in rp_patterns:
        rp_matches.extend(glob.glob(pat))
    rp_matches = sorted(set(rp_matches))

    # ③ 扫描工作流历史
    wf_dir = "D:/TestScheduler/memory/workflow_states"
    wf_matches = []
    if os.path.isdir(wf_dir):
        for fname in os.listdir(wf_dir):
            if not fname.endswith(".json"):
                continue
            if req_id in fname:
                wf_matches.append(os.path.join(wf_dir, fname))
    wf_matches = sorted(wf_matches)

    exists = bool(tc_matches or rp_matches or wf_matches)

    return json.dumps({
        "ok": True,
        "exists": exists,
        "requirement_id": req_id,
        "files": tc_matches,
        "reports": rp_matches,
        "workflows": wf_matches,
        "message": (
            f"检测到需求 {req_id} 的已有产物："
            f"用例文件 {len(tc_matches)} 个，"
            f"执行报告 {len(rp_matches)} 个，"
            f"工作流记录 {len(wf_matches)} 个"
            if exists else
            f"需求 {req_id} 暂无历史用例、报告或工作流记录"
        ),
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════
# 8. check-existing-report
# ═══════════════════════════════════════════════════════

def check_existing_report(params: dict) -> str:
    """检查需求是否已有一份完整的测试执行报告，返回结构化结果"""
    req_id = params.get("requirement_id", "").strip()
    if not req_id:
        return json.dumps({"ok": False, "error": "缺少 requirement_id"}, ensure_ascii=False)

    import glob

    # 报告文件命名规则: {req_id}_report_*.html 或 {req_id}_report_*.json
    patterns = [
        os.path.join(REPORT_DIR, f"{req_id}_report_*.html"),
        os.path.join(REPORT_DIR, f"{req_id}_report_*.json"),
    ]

    all_matches = []
    for pat in patterns:
        all_matches.extend(glob.glob(pat))

    if all_matches:
        files = sorted(all_matches)
        return json.dumps({
            "ok": True,
            "exists": True,
            "requirement_id": req_id,
            "files": files,
            "message": f"检测到需求 {req_id} 已有 {len(files)} 份测试报告",
        }, ensure_ascii=False, indent=2)

    return json.dumps({
        "ok": True,
        "exists": False,
        "requirement_id": req_id,
        "files": [],
        "message": f"需求 {req_id} 暂无历史测试报告",
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════
# 9. 主入口
# ═══════════════════════════════════════════════════════

def run(params: dict, context: dict | None = None) -> str:
    action = params.get("action", "").strip().lower()

    if action == "init-materials":
        return init_materials(params)
    elif action == "generate-excel":
        return generate_excel(params)
    elif action == "convert-interface-doc":
        return convert_interface_doc(params)
    elif action == "list-channels":
        return list_channels(params)
    elif action == "list-interface-docs":
        return list_interface_docs(params)
    elif action == "find-interface-doc":
        return find_interface_doc(params)
    elif action == "scan-materials":
        return scan_materials(params)
    elif action == "check-existing-testcases":
        return check_existing_testcases(params)
    elif action == "check-existing-report":
        return check_existing_report(params)
    else:
        return json.dumps({
            "ok": False,
            "error": f"未知操作: {action}",
            "usage": {
                "init-materials": "创建需求材料目录和模板。参数: requirement_id",
                "scan-materials": "扫描需求材料目录。参数: requirement_id",
                "generate-excel": "生成用例Excel。参数: cases_json, output_name",
                "convert-interface-doc": "转换接口文档为md。参数: source_file, channel_name",
                "find-interface-doc": "根据渠道名查找接口文档(优先MD缓存)。参数: channel_name",
                "list-channels": "列出可用渠道JSON",
                "list-interface-docs": "列出可用接口文档(含缓存MD)",
                "check-existing-testcases": "检查是否已有用例文件。参数: requirement_id",
                "check-existing-report": "检查是否已有测试报告。参数: requirement_id",
            },
        }, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "error": "用法: run.py [action] [json_params]"}, ensure_ascii=False))
        sys.exit(1)

    action = sys.argv[1]
    if len(sys.argv) >= 3:
        try:
            input_params = json.loads(sys.argv[2])
        except json.JSONDecodeError:
            print(json.dumps({"ok": False, "error": "参数必须是有效 JSON"}, ensure_ascii=False))
            sys.exit(1)
    else:
        input_params = {}

    input_params["action"] = action
    result = run(input_params)
    print(result)
